from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_urfu_plan(data, filename="Individual_Plan.docx"):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # 1. ШАПКА ДОКУМЕНТА

    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.add_run("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ\n")
    header_p.add_run("Федеральное государственное автономное образовательное учреждение высшего образования\n")
    header_p.add_run("«Уральский федеральный университет имени первого Президента России Б.Н.Ельцина»")

    # 2. БЛОК "УТВЕРЖДАЮ"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.add_run("УТВЕРЖДАЮ\n").bold = True
    p.add_run("Руководитель образовательной программы\n")
    p.add_run("__________ ").bold = True
    p.add_run(f"{data['rop_name']}\n")
    p.add_run(f"«___» _________ {data['year']} г.")

    doc.add_paragraph()

    # 3. ЗАГОЛОВОК И ФИО СТУДЕНТА

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run("ИНДИВИДУАЛЬНЫЙ УЧЕБНЫЙ ПЛАН МАГИСТРАНТА").bold = True

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.add_run(data["student_name"]).bold = True

    # 4. ОБЩИЕ СВЕДЕНИЯ (ТАБЛИЦА)

    doc.add_paragraph("Общие сведения").bold = True

    table_info = doc.add_table(rows=0, cols=2)
    table_info.style = "Table Grid"  # Включаем границы

    def add_info_row(label, value):
        row = table_info.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    add_info_row("Институт", data["institute"])
    add_info_row("Кафедра/Департамент", data["department"])
    add_info_row("Направление подготовки", data["specialty_code"])
    add_info_row("Образовательная программа", data["program"])
    add_info_row("Научный руководитель", data["supervisor"])
    add_info_row("Направление научно-исследовательской деятельности", data["research_area"])
    add_info_row("Тема выпускной квалификационной работы", data["thesis_topic"])
    add_info_row("Срок предоставления ВКР к защите", data["deadline"])

    doc.add_paragraph()

    # 5. СЕМЕСТРЫ И ДИСЦИПЛИНЫ

    def create_subject_table(subjects_list):
        if not subjects_list:
            doc.add_paragraph("Дисциплины не выбраны").italic = True
            return

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"

        headers = [
            "№ п/п",
            "Наименование дисциплин, практик",
            "Объем аудит. работы, час",
            "Объем (зет)",
            "Форма аттестации",
            "Оценка",
        ]
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, subj in enumerate(subjects_list, 1):
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = subj["name"]
            row_cells[2].text = str(subj["hours"])
            row_cells[3].text = str(subj["credits"])
            row_cells[4].text = subj["control"]
            row_cells[5].text = ""

    for semester in data["semesters"]:
        doc.add_paragraph()

        p_sem = doc.add_paragraph()
        p_sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sem.add_run(f"{semester['number']} семестр обучения, {semester['academic_year']} гг.").bold = True

        doc.add_paragraph("Образовательная часть программы подготовки")
        doc.add_paragraph("Обязательная часть в соответствии с учебным планом").italic = True

        create_subject_table(semester.get("mandatory_subjects", []))

        doc.add_paragraph()
        doc.add_paragraph("Дисциплины и курсы по выбору и факультативы").italic = True

        create_subject_table(semester.get("elective_subjects", []))

        doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6. ПОДПИСИ
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.add_run(f"Магистрант {data['student_short_name']} ______________")

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"«___» _________ {data['year']} г.")

    p_rop = doc.add_paragraph()
    p_rop.add_run(f"Руководитель магистерской программы (РОП) {data['rop_name']} ______________")

    p_sci = doc.add_paragraph()
    p_sci.add_run("Научный руководитель ________________________")

    doc.add_paragraph("Заключение кафедры:")
    doc.add_paragraph("_" * 80)

    p_prot = doc.add_paragraph()
    p_prot.add_run("Протокол № ____ от «___» _________ 20___ г.")

    doc.save(filename)
    print(f"Файл {filename} успешно создан.")  # noqa: T201


# ДАННЫЕ ДЛЯ ТЕСТА
sample_data = {
    "student_name": "Гилемшина Рината Ришатовича",
    "student_short_name": "Р.Р. Гилемшин",
    "rop_name": "А.А. Кошелев",
    "year": "2025",
    "institute": "Естественных наук и математики",
    "department": "Школа наук",
    "specialty_code": "02.04.01 Математика и компьютерные науки",
    "program": "Современные проблемы компьютерных наук",
    "supervisor": "",
    "research_area": "",
    "thesis_topic": "",
    "deadline": "Май 2027",
    "semesters": [
        {
            "number": 1,
            "academic_year": "2025 / 2026",
            "mandatory_subjects": [
                {"name": "Современные научные исследования", "hours": 34, "credits": 4, "control": "Экзамен"},
                {"name": "Компьютерные науки", "hours": 34, "credits": 2, "control": "Зачет"},
                {"name": "Современные задачи компьютерных наук", "hours": 34, "credits": 6, "control": "Экзамен"},
            ],
            "elective_subjects": [
                {"name": "Введение в Python", "hours": 34, "credits": 3, "control": "Зачет"},
                {"name": "Введение в Kaggle", "hours": 34, "credits": 3, "control": "Зачет"},
            ],
        },
        {
            "number": 2,
            "academic_year": "2025 / 2026",
            "mandatory_subjects": [
                {"name": "Иностранный язык в профессиональной сфере", "hours": 34, "credits": 3, "control": "Зачет"},
                {"name": "Управление IT-проектами", "hours": 34, "credits": 4, "control": "Экзамен"},
            ],
            "elective_subjects": [
                {"name": "Глубокое обучение (Deep Learning)", "hours": 34, "credits": 5, "control": "Экзамен"},
                {"name": "DevOps практики", "hours": 34, "credits": 3, "control": "Зачет"},
            ],
        },
    ],
}
if __name__ == "__main__":
    create_urfu_plan(sample_data)
